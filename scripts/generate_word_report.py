#!/usr/bin/env python3
from __future__ import annotations

import json
from pathlib import Path
from typing import List

import pandas as pd


def load_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path) if path.exists() else pd.DataFrame()


def add_heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc, items: List[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_image(doc, path: Path, width_inches: float = 6.5) -> None:
    from docx.shared import Inches
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_inches))


def add_table_from_df(doc, df: pd.DataFrame, max_rows: int | None = None) -> None:
    if df.empty:
        add_paragraph(doc, "No data available.")
        return
    if max_rows is not None:
        df = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row.values):
            cells[i].text = str(val)


def tuned_vs_base(lines: List[str], best_name: str, results: dict, label: str) -> List[str]:
    if isinstance(best_name, str) and best_name.startswith("Tuned "):
        base_name = best_name.replace("Tuned ", "", 1)
        tuned = results.get(best_name, {}) or {}
        base = results.get(base_name, {}) or {}
        try:
            t_mae = float(tuned.get("test_mae"))
            t_r2 = float(tuned.get("test_r2"))
            b_mae = float(base.get("test_mae"))
            b_r2 = float(base.get("test_r2"))
            d_mae = b_mae - t_mae
            d_r2 = t_r2 - b_r2
            lines.append(f"Tuned vs Base – {label}")
            lines.append(f"  • Base: {base_name} — MAE {b_mae:.2f}, R² {b_r2:.4f}")
            lines.append(f"  • Tuned: {best_name} — MAE {t_mae:.2f}, R² {t_r2:.4f}")
            lines.append(f"  • Improvement: ΔMAE {d_mae:+.2f} (lower is better), ΔR² {d_r2:+.4f}")
        except Exception:
            pass
    return lines


def main() -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("Please install python-docx: pip install python-docx")
        raise

    repo_root = Path(__file__).resolve().parents[1]
    out_dir = repo_root / "output"
    rep_dir = repo_root / "reports"
    rep_dir.mkdir(parents=True, exist_ok=True)

    # Load artifacts
    results_path = out_dir / "enhanced_analysis_results.json"
    ev_rank_path = out_dir / "enhanced_ev_model_rankings.csv"
    ice_rank_path = out_dir / "enhanced_ice_model_rankings.csv"

    results = json.loads(results_path.read_text(encoding="utf-8")) if results_path.exists() else {}
    ev_df = load_csv(ev_rank_path)
    ice_df = load_csv(ice_rank_path)

    best_ev = results.get("best_ev_model", "")
    best_ice = results.get("best_ice_model", "")
    ev_results = results.get("ev_results", {})
    ice_results = results.get("ice_results", {})
    perf = results.get("performance_summary", {})
    specs = results.get("system_specs", {})

    doc = Document()

    # Title
    add_heading(doc, "Enhanced Vehicle Efficiency Analysis – Technical Report", level=0)
    add_paragraph(doc, "This report documents the data, methods, modeling, fine-tuning, and results for predicting vehicle efficiency across EV and ICE cohorts.")

    # Abstract
    add_heading(doc, "Abstract", level=1)
    add_paragraph(doc, "We analyze a vehicle dataset, engineer domain-informed features, select predictive subsets, train and evaluate multiple regressors, and fine-tune the top candidates. The pipeline outputs ranked results, serialized models, and comprehensive visualizations. This document explains the process and summarizes results, including any tuned-vs-base improvements.")

    # Data
    add_heading(doc, "Data & Preprocessing", level=1)
    add_bullets(doc, [
        "Input CSV: data/vehicle_comparison_dataset_030417.csv",
        "Target: efficiency = mileage_km / energy_consumption (computed)",
        "Outlier removal: IQR filter on efficiency",
        "Cohorts: split by vehicle_type → EV and ICE; drop vehicle_type thereafter",
        "Clipping: co2_emissions_g_per_km clipped at lower=0",
    ])

    # Feature Engineering
    add_heading(doc, "Feature Engineering & Selection", level=1)
    add_bullets(doc, [
        "Engineered ratios: power_efficiency, storage_per_torque, cost_efficiency",
        "Maintenance & lifespan: maintenance_per_year, maintenance_per_torque, lifespan_torque_ratio",
        "Environmental: eco_efficiency (EV), green_performance (EV), emission_intensity (ICE), emission_per_storage (ICE)",
        "Categories → codes: torque_category_num, acceleration_category_num",
        "Polynomial/interactions: torque_squared, cost_squared, torque_x_lifespan, storage_x_lifespan",
        "Transforms/normalization: log_* features and normalized_* features",
        "Leakage prevention: exclude mileage_km and energy_consumption from features",
        "Selection: correlation thresholds (EV>0.02, ICE>0.03), variance filter, multicollinearity pruning (|corr|<0.85), cap 20 features",
    ])

    # Modeling
    add_heading(doc, "Modeling & Evaluation", level=1)
    add_bullets(doc, [
        "Models: Linear, Ridge, Lasso, Random Forest, Gradient Boosting, Decision Tree; optional XGBoost, LightGBM, CatBoost",
        "Preprocessing: PowerTransformer for linear models; raw passthrough for tree/boosting",
        "Validation: 5-fold KFold (shuffle, random_state=42), metric: MAE",
        "Hold-out: 80/20 train/test split; metrics: MAE, RMSE, R²",
        "Ranking: configurable by R² or MAE (default R²)",
    ])

    # Fine-tuning
    add_heading(doc, "Fine-Tuning Strategy", level=1)
    add_bullets(doc, [
        "Select top 2 models per cohort (by chosen rank metric)",
        "Boosting models: Optuna search if available (trials configurable), fallback to sklearn",
        "Non-boosting: RandomizedSearchCV by default, or GridSearchCV",
        "Tuning metric: MAE or R² (default MAE)",
    ])

    # Results summary
    add_heading(doc, "Results Summary", level=1)
    if perf:
        add_paragraph(doc, f"Best EV Model: {best_ev} (Test R²: {perf.get('best_ev_r2')})")
        add_paragraph(doc, f"Best ICE Model: {best_ice} (Test R²: {perf.get('best_ice_r2')})")
    else:
        add_paragraph(doc, f"Best EV Model: {best_ev}")
        add_paragraph(doc, f"Best ICE Model: {best_ice}")

    # Tuned vs Base (if applicable)
    snippet_lines: List[str] = []
    tuned_vs_base(snippet_lines, best_ev, ev_results, "EV")
    tuned_vs_base(snippet_lines, best_ice, ice_results, "ICE")
    if snippet_lines:
        add_heading(doc, "Tuned vs Base", level=2)
        for ln in snippet_lines:
            add_paragraph(doc, ln)

    # Rankings tables
    add_heading(doc, "EV Model Rankings (Top 5)", level=2)
    add_table_from_df(doc, ev_df, max_rows=5)
    add_heading(doc, "ICE Model Rankings (Top 5)", level=2)
    add_table_from_df(doc, ice_df, max_rows=5)

    # Visualizations
    add_heading(doc, "Visualizations", level=1)
    add_paragraph(doc, "The following figures summarize distributions, model performance, and correlation structures.")
    add_heading(doc, "Main Dashboard", level=2)
    add_image(doc, out_dir / "enhanced_efficiency_dashboard.png")
    add_heading(doc, "CV Stability", level=2)
    add_image(doc, out_dir / "cv_stability_comparison.png", width_inches=5.5)
    add_heading(doc, "Correlation Dashboards", level=2)
    add_image(doc, out_dir / "correlation_analysis_dashboard.png")
    add_image(doc, out_dir / "detailed_correlation_matrices.png")

    # System specs
    if specs:
        add_heading(doc, "System & Reproducibility", level=1)
        for k, v in specs.items():
            add_paragraph(doc, f"- {k}: {v}")
        add_paragraph(doc, "Random state used in KFold and model constructors where applicable.")

    # Save document
    out_doc = rep_dir / "final_modeling_report.docx"
    doc.save(str(out_doc))
    print(f"Word report written to {out_doc}")


if __name__ == "__main__":
    main()

